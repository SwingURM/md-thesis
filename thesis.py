import os
from collections import deque

from docx import Document

from header import (
    add_toc,
    fix_reference_format,
    force_update_fields,
    insert_section_breaks,
    pageBreak_before,
    process_hyperlink,
    process_math_equations,
    process_table,
    replace_ref_format_in_doc,
    set_all_secs_thesis,
    set_headers_thesis,
    update_reference_doc, set_abstract_font,
)

TITLE = "面向优秀论文标准的研究"


# --- Main Processing ---
def process_document(doc):
    """Process the Word document."""

    deque(map(process_table, doc.tables))

    add_toc(doc)

    deque(map(pageBreak_before, doc.paragraphs))

    insert_section_breaks(doc, 5)

    set_all_secs_thesis(doc)

    set_headers_thesis(doc)

    process_math_equations(doc)

    set_abstract_font(doc)

    process_hyperlink(doc)

    replace_ref_format_in_doc(doc)

    force_update_fields(doc)

    fix_reference_format(doc)


# --- Entry Point ---
if __name__ == "__main__":
    INPUT = "demo.md"
    REF_FILE = "cppref.bib"
    output = f"{TITLE}-generated.docx"

    update_reference_doc()
    assert (
        os.system(
            f"pandoc {INPUT} -o {output} --filter pandoc-crossref --reference-doc reference.docx --citeproc --csl GB-T-7714—2015（顺序编码，双语，姓名不大写，无URL、DOI，引注有页码）.csl --bibliography {REF_FILE}"
        )
        == 0
    ), "pandoc execution failed"

    doc = Document(output)
    assert doc is not None, "Failed to load the document"
    process_document(doc)
    doc.save(output)
    print("Output file saved as:", output)
