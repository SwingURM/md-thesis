import os
import re
import shutil
from collections import deque

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

title = "面向优秀论文标准的研究"
statistics_data = {}


def update_reference_doc():
    shutil.make_archive("reference", "zip", "reference")
    if os.path.exists("reference.docx"):
        os.remove("reference.docx")
    os.rename("reference.zip", "reference.docx")


# --- Helper Functions ---
def create_element(name):
    """Create an OXML element."""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """Set an attribute for an OXML element."""
    element.set(qn(name), value)


def delete_paragraph(paragraph):
    """Remove a paragraph element."""
    p = paragraph._element
    if p is not None and p.getparent() is not None:
        p.getparent().remove(p)


def add_page_number_field(paragraph):
    """Add a PAGE field to a paragraph."""
    run = paragraph.add_run()
    fldSimple = create_element("w:fldSimple")
    create_attribute(fldSimple, "w:instr", r" PAGE \* MERGEFORMAT ")
    run._r.append(fldSimple)


def set_page_number_style(section, fmt="decimal", start=None):
    """Set page number format and starting number for a section."""
    sectPr = section._sectPr
    assert sectPr.find(qn("w:pgNumType")) is None, "Section already has pgNumType"
    pgNumType = create_element("w:pgNumType")
    create_attribute(pgNumType, "w:fmt", fmt)
    if start is not None:
        create_attribute(pgNumType, "w:start", str(start))
    sectPr.append(pgNumType)


def apply_simsun_tnr_font(run):
    """Apply SimSun font to a run."""
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def copy_sectPr_properties(source_sectPr, target_sectPr):
    """Copy key page settings from source sectPr to target_sectPr."""
    properties_to_copy = ["pgSz", "pgMar", "cols", "docGrid"]

    for prop_tag in properties_to_copy:
        source_element = source_sectPr.find(qn(f"w:{prop_tag}"))
        if source_element is not None:
            # Create a new element with the same tag
            target_element = create_element(f"w:{prop_tag}")

            # Copy all attributes
            for name, value in source_element.attrib.items():
                ns, localname = name.split("}") if "}" in name else ("", name)
                localname = localname.split(":")[-1] if ":" in localname else localname
                create_attribute(target_element, f"w:{localname}", value)

            # Replace existing or append
            existing = target_sectPr.find(qn(f"w:{prop_tag}"))
            if existing is not None:
                target_sectPr.replace(existing, target_element)
            else:
                target_sectPr.append(target_element)


def add_next_page_section_break(paragraph):
    """Add a next-page section break to a paragraph."""
    assert paragraph is not None, "Paragraph cannot be None"
    p_element = paragraph._p
    pPr = p_element.get_or_add_pPr()

    sectPr = pPr.find(qn("w:sectPr"))
    if sectPr is None:
        sectPr = create_element("w:sectPr")
        pPr.append(sectPr)
        doc = paragraph._parent.part.document
        last_section_sectPr = doc.sections[-1]._sectPr
        copy_sectPr_properties(last_section_sectPr, sectPr)

    # Set section break type to next page
    type_element = sectPr.find(qn("w:type"))
    if type_element is None:
        type_element = create_element("w:type")
        sectPr.append(type_element)
    create_attribute(type_element, "w:val", "nextPage")

    # Clear paragraph content (as it's just a marker)
    paragraph.text = ""
    for run in paragraph.runs:
        p_element.remove(run._r)


def para_is_style(paragraph, style_name):
    """Check if a paragraph is an abstract paragraph."""
    return paragraph.style.name.lower() == style_name.lower()


def add_toc(document):
    """
    Insert a Table of Contents (TOC) at the location of a marker text.
    """
    # find all the paragraph with style toc Heading
    target_paragraphs = [
        p for p in document.paragraphs if para_is_style(p, "toc heading")
    ]
    assert len(target_paragraphs) == 1, (
        "The document must contain exactly one paragraph with the style 'TOC Heading'."
    )
    target_paragraph = target_paragraphs[0]
    try:
        toc_title_paragraph = target_paragraph.insert_paragraph_before(
            "目录", style="TOC Heading"
        )
        toc_field_paragraph = target_paragraph.insert_paragraph_before("")
        run = toc_field_paragraph.add_run()

        fldChar_begin = create_element("w:fldChar")
        create_attribute(fldChar_begin, "w:fldCharType", "begin")

        instrText = create_element("w:instrText")
        create_attribute(instrText, "xml:space", "preserve")
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

        fldChar_end = create_element("w:fldChar")
        create_attribute(fldChar_end, "w:fldCharType", "end")

        run._r.extend([fldChar_begin, instrText, fldChar_end])
        delete_paragraph(target_paragraph)
    except Exception as e:
        print(f"Error during TOC insertion: {e}")


def insert_section_breaks(doc, num):
    """Insert section breaks at specified markers."""

    # find paragraphs with style "myBreak"
    break_paragraphs = [p for p in doc.paragraphs if para_is_style(p, "pBreak")]
    assert len(break_paragraphs) == num, (
        f"The document must contain exactly {num} paragraphs with the style 'pBreak', found {len(break_paragraphs)}.")
    deque(map(add_next_page_section_break, break_paragraphs))

def pageBreak_before(paragraph):
    """Set page break before and w:after for each heading 1."""
    """this is for thesis only"""
    if not para_is_style(paragraph, "heading 1"):
        return
    pPr = paragraph._element.get_or_add_pPr()
    pageBreak = create_element("w:pageBreakBefore")
    create_attribute(pageBreak, "w:val", "true")
    w_after = create_element("w:spacing")
    create_attribute(w_after, "w:after", "240")
    pPr.append(w_after)
    pPr.append(pageBreak)



def process_table(table):
    """Apply formatting to a table."""
    tblPr = table._element.xpath(".//w:tblPr")[0]

    # Set table width to auto
    tblW = create_element("w:tblW")
    create_attribute(tblW, "w:w", "0")
    create_attribute(tblW, "w:type", "auto")
    tblPr.append(tblW)

    # Set table borders
    tblBorders = create_element("w:tblBorders")
    for border in ["top", "bottom"]:
        borderElement = create_element(f"w:{border}")
        create_attribute(borderElement, "w:val", "single")
        create_attribute(borderElement, "w:sz", "12")
        create_attribute(borderElement, "w:space", "0")
        create_attribute(borderElement, "w:color", "000000")
        tblBorders.append(borderElement)
    tblPr.append(tblBorders)

    # Set table style
    tblLook = create_element("w:tblLook")
    create_attribute(tblLook, "w:val", "04A0")
    create_attribute(tblLook, "w:firstRow", "1")
    create_attribute(tblLook, "w:lastRow", "0")
    create_attribute(tblLook, "w:firstColumn", "1")
    create_attribute(tblLook, "w:lastColumn", "0")
    create_attribute(tblLook, "w:noHBand", "0")
    create_attribute(tblLook, "w:noVBand", "1")
    tblPr.append(tblLook)

    # Center align the table
    tblAlignment = create_element("w:jc")
    create_attribute(tblAlignment, "w:val", "center")
    tblPr.append(tblAlignment)

    # Calculate column width
    TOTAL_WIDTH_TWIPS = 9286  # Total table width
    num_columns = len(table.columns)  # Number of columns
    column_width_twips = TOTAL_WIDTH_TWIPS // num_columns  # Column width

    # Apply style to the first row
    first_row = table.rows[0]  # Get the first row

    def set_cell_style(cell):
        tcPr = cell._element.get_or_add_tcPr()  # Get or create cell properties
        # Set bottom border
        tcBorders = create_element("w:tcBorders")
        bottomBorder = create_element("w:bottom")
        create_attribute(bottomBorder, "w:val", "single")
        create_attribute(bottomBorder, "w:sz", "6")
        create_attribute(bottomBorder, "w:space", "0")
        create_attribute(bottomBorder, "w:color", "000000")
        tcBorders.append(bottomBorder)
        tcPr.append(tcBorders)

    deque(map(set_cell_style, first_row.cells))

    for column in table.columns:
        for cell in column.cells:
            cell.width = Pt(column_width_twips / 20)  # Convert to points


def set_all_secs_thesis(doc):
    """Apply formatting to document sections."""
    num_sections = len(doc.sections)
    print(f"Document contains {num_sections} sections.")

    assert num_sections == 6, "Document must have at least 6 sections."

    # Section 1,2: Before TOC
    section1 = doc.sections[0]
    section1.footer_distance = Pt(56.7)  # Footer distance from the bottom (2 cm)
    set_page_number_style(section1, fmt="upperRoman", start=1)
    add_page_number_to_footer(section1)

    section2 = doc.sections[1]
    section2.footer_distance = Pt(56.7)  # Footer distance from the bottom (2 cm)
    set_page_number_style(section2, fmt="upperRoman")

    # Section 3: TOC
    section3 = doc.sections[2]
    section3.footer_distance = Pt(56.7)  # Footer distance from the bottom (2 cm)
    section3.footer.is_linked_to_previous = False
    set_page_number_style(section3, fmt="upperRoman", start=1)
    add_page_number_to_footer(section3)

    # Section 4: After TOC
    section4 = doc.sections[3]
    section4.footer.is_linked_to_previous = False
    set_page_number_style(section4, fmt="decimal", start=1)
    for i in range(3, 6):
        section = doc.sections[i]
        section.footer_distance = Pt(56.7)  # Footer distance from the bottom (2 cm)


def set_all_secs_other(doc):
    """Apply formatting to document sections."""
    num_sections = len(doc.sections)
    print(f"Document contains {num_sections} sections.")

    assert num_sections == 2, "Document must have 2 sections."

    section1 = doc.sections[0]
    set_page_number_style(section1, fmt="decimal", start=1)

def set_headers_other(doc):
    """
    Set headers for the only sections in the document.
    """
    assert len(doc.sections) == 2, "Document must have exactly 2 section."
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    # Get the header of the section
    header = section.header
    assert len(header.paragraphs) == 1, "Header must contain one paragraph"
    paragraph = header.paragraphs[0]
    # Clear existing content in the paragraph
    for run in paragraph.runs:
        run.clear()
    # Add header content
    run = paragraph.add_run(title + " ")
    run_append_page_number(run)
    # Set paragraph style
    paragraph.style = "header"  # Ensure the style name is correct
    # Set header distance
    section.header_distance = Pt(56.7)  # Header distance from the top (2 cm)
    

def run_append_page_number(run):
    # Add page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")  # Avoid truncation of field instructions
    instrText.text = "PAGE \\* MERGEFORMAT"
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)


def set_headers_thesis(doc):
    """
    Set headers for all sections in the document.
    - First 3 sections: Header without page numbers.
    - Following section: Header with page numbers.
    """
    for i, section in enumerate(doc.sections):
        section.header.is_linked_to_previous = False
        # Get the header of the section
        header = section.header

        assert len(header.paragraphs) == 1, "Header must contain one paragraph"
        paragraph = header.paragraphs[0]

        # Clear existing content in the paragraph
        for run in paragraph.runs:
            run.clear()

        # Add header content
        if i < 3:  # First 3 sections
            paragraph.add_run(title + " ")
        else:  # Fourth section and beyond
            run = paragraph.add_run(title + " ")
            run_append_page_number(run)

        # Set paragraph style
        paragraph.style = "header"  # Ensure the style name is correct

        # Set header distance
        section.header_distance = Pt(56.7)  # Header distance from the top (2 cm)


def add_page_number_to_footer(section):
    """Add page number to the footer of a section."""
    footer = section.footer
    assert len(footer.paragraphs) == 1, "Footer must contain a existing paragraph"
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.style = "footer"
    add_page_number_field(paragraph)

def split_by_colon(text):
    assert ':' in text or '：' in text, (
        "Abstract paragraph must contain '：' or ':' to separate prefix and content")
    # 尝试用英文冒号分割
    if ':' in text:
        parts = text.split(':', 1)  # 只分割第一个冒号
        return parts[0].strip() + ':', parts[1].strip()
    
    # 尝试用中文冒号分割
    if '：' in text:
        parts = text.split('：', 1)  # 只分割第一个冒号
        return parts[0].strip() + "：", parts[1].strip()

def set_abstract_font(doc):
    """Set the font for all abstract paragraphs."""
    print("\n=== Processing abstract paragraphs ===")
    for paragraph in doc.paragraphs:
        if para_is_style(paragraph, "abstract"):
            text = paragraph.text.strip()
            part1, part2 = split_by_colon(text)

            print(f"Processing '{part1}...'")
            paragraph.clear()

            # Add prefix with bold formatting
            run_prefix = paragraph.add_run(part1)
            if part1 == "Keywords:":
                run_prefix.bold = True  # Apply bold to the prefix

            # Add the remaining text with appropriate font
            run_rest = paragraph.add_run(part2)
            apply_simsun_tnr_font(run_rest)

    print("=== Completed processing abstract paragraphs ===\n")


def replace_ref_format_in_doc(doc):
    """
    Replace all occurrences of '图%d.%d' with '图%d-%d' in the entire Word document
    while preserving oMath elements and other formatting.
    """
    figure_pattern = re.compile(r"图(\d+)\.(\d+)")
    eqn_pattern = re.compile(r"式(\d+)\.(\d+)")

    for paragraph in doc.paragraphs:
        # Process the paragraph run by run to preserve formatting
        for run in paragraph.runs:
            if run.text and ("图" in run.text or "式" in run.text):
                run.text = figure_pattern.sub(r"图\1-\2", run.text)
                run.text = eqn_pattern.sub(r"式\1-\2", run.text)


def process_math_equations(doc):
    """
    Locate and format all oMathParagraph elements in the document.
    - Add right-aligned tab stop
    - Structure with equation, tab, and number
    - Detect and extract equation numbers in format (d.d)
    """
    print("\n=== STARTING MATH PARAGRAPH DETECTION AND FORMATTING ===")
    math_para_count = 0

    for i, paragraph in enumerate(doc.paragraphs):
        # Check if paragraph contains math paragraph elements
        if "<m:oMathPara" in paragraph._p.xml:
            math_para_count += 1
            # print(f"\n--- Math Paragraph #{math_para_count} found in paragraph {i} ---")
            all_math_t = paragraph._element.xpath(".//m:t")

            assert len(all_math_t) >= 3, "数学段落必须包含至少3个数学文本节点"
            last_4_text = all_math_t[-4].text
            assert last_4_text.strip("\u2001\u2003\u2000\u2002 ") == "", (
                "数学段落的倒数第四个数学文本节点必须是空格"
            )
            last_3_texts = [t.text for t in all_math_t[-3:]]
            assert last_3_texts[0] == "(" and last_3_texts[-1] == ")", (
                "数学段落的最后三个数学文本节点必须中有'('和')'"
            )
            equation_number = last_3_texts[1]
            assert re.match(r"^\d+\.\d+$", equation_number), (
                "数学段落的倒数第三个数学文本节点必须是数字编号，格式为d.d"
            )
            equation_number = equation_number.replace(".", "-")

            # Remove the last 4 m:t nodes
            for t_node in all_math_t[-4:]:
                r_node = t_node.getparent()

                if r_node is not None:
                    parent = r_node.getparent()
                    if parent is not None:
                        parent.remove(r_node)
            # Format the paragraph with oMathPara, passing the discovered number
            format_math_paragraph(paragraph, equation_number)

    print(
        f"\n=== COMPLETED MATH ELEMENT FORMATTING: {math_para_count} oMathPara elements formatted ===\n"
    )


def format_math_paragraph(paragraph, equation_number="replace_me"):
    """
    Format a paragraph containing oMathPara with proper style and equation numbering.

    Args:
        paragraph: The paragraph containing the math element
        equation_number: The sequential number to assign to this equation
    """
    try:
        # Set paragraph style to FormulaEquationNumbered
        paragraph.style = "FormulaEquationNumbered"

        # Add the equation number in parentheses after a tab
        # First get the XML element
        p_xml = paragraph._element

        # Create a run element for the tab and equation number
        run_xml = create_element("w:r")

        # Add a tab character
        tab_xml = create_element("w:tab")
        run_xml.append(tab_xml)

        # Add the equation number in parentheses
        text_xml = create_element("w:t")
        text_xml.text = f"（{equation_number}）"
        run_xml.append(text_xml)

        # Append the new run to the paragraph
        p_xml.append(run_xml)

        # print(
        #     f"Applied 'FormulaEquationNumbered' style with equation number {equation_number}"
        # )
        return True
    except Exception as e:
        print(f"Error formatting math paragraph: {e}")
        return False


def process_hyperlink(doc):
    """
    取消文档中非上标的所有超链接,上标的超链接设置特定样式
    """
    print("\n=== PROCESSING HYPERLINKS ===")
    count1 = 0
    count2 = 0

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        p_xml = paragraph._element

        for hyperlink in p_xml.xpath(".//w:hyperlink"):
            assert len(hyperlink.xpath(".//w:r")) == 1, (
                "Hyperlink must contain exactly one run"
            )
            r = hyperlink.xpath(".//w:r")[0]
            parent = hyperlink.getparent()

            assert r.tag.endswith("r"), "Child of hyperlink must be a run"
            assert r.rPr is not None, "Run must have rPr"
            vertAlign = r.rPr.xpath(".//w:vertAlign[@w:val='superscript']")
            if vertAlign:
                count2 += 1
                text_element = r.xpath(".//w:t")
                assert len(text_element) == 1, (
                    "Run must contain exactly one text element"
                )
                text_content = text_element[0].text
                style_element = create_element("w:rStyle")
                create_attribute(style_element, "w:val", "ae")
                assert r.rPr.rStyle.val == "af", (
                    "Run must have style 'af' before replacing"
                )
                r.rPr.replace(r.rPr.rStyle, style_element)
                statistics_data[text_content] = statistics_data.get(text_content, 0) + 1
            else:
                count1 += 1
                r.remove(r.rPr)
                parent.insert(parent.index(hyperlink), r)

                parent.remove(hyperlink)

    print(
        f"Total hyperlinks removed: {count1 + count2}, Non-superscript hyperlinks removed: {count1}, Superscript hyperlinks retained: {count2}"
    )


def has_chinese(text):
    return bool(re.search(r"[\u4e00-\u9fff]", text.replace("等", "")))


def fix_reference_format(doc):
    print("\n=== FIXING REFERENCE FORMATTING ===")
    is_reference_section = False
    ref_ch = 0
    ref_en = 0
    ref_fixed = 0
    for paragraph in doc.paragraphs:
        if (
            para_is_style(paragraph, "heading 1")
            and paragraph.text.strip() == "参考文献"
        ):
            is_reference_section = True
            continue
        if not is_reference_section:
            continue
        if is_reference_section and para_is_style(paragraph, "pBreak"):
            return
        assert paragraph.runs and len(paragraph.runs) >= 1, (
            "Reference paragraph must have at least one run"
        )
        run1 = paragraph.runs[0]
        if not run1.text.startswith("["):
            continue
        if has_chinese(paragraph.text):
            ref_ch += 1
            continue
        ref_en += 1

        def etal_replace(run):
            nonlocal ref_fixed
            if run.text == "等.":
                run.text = "et al."
                ref_fixed += 1

        deque(map(etal_replace, paragraph.runs))
    print(
        f"Total references processed: {ref_ch + ref_en}, Chinese references: {ref_ch}, English references: {ref_en}, '等.' replaced with 'et al.': {ref_fixed}"
    )


def force_update_fields(doc):
    """强制 Word 在打开时提示更新域（包括目录）"""
    element = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    element.append(update_fields)
